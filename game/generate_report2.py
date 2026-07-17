import docx

doc = docx.Document('D:/嫣落尘渊/孙若妍生产实习报告.docx')

def replace_paragraph(para, new_text):
    para.text = new_text

# 找到2.3系统功能设计的位置
for i, para in enumerate(doc.paragraphs):
    if para.text.strip() == '2.3 系统功能设计':
        section_start = i
        break

# 找到3模块设计的位置
section_end = len(doc.paragraphs)
for i, para in enumerate(doc.paragraphs):
    if '3 模块设计' in para.text.strip():
        section_end = i
        break

# 清空2.3节到3节之间的内容
for i in range(section_start + 1, section_end):
    doc.paragraphs[i].text = ''

# 更新2.3节标题和内容
replace_paragraph(doc.paragraphs[section_start], "2.3 系统功能设计")
replace_paragraph(doc.paragraphs[section_start + 1], "本系统的功能设计以玩家体验为核心，深度融合古风穿越题材特性与现代Web技术，旨在构建一个交互丰富、视觉精美的沉浸式文字冒险游戏。通过对乙女游戏用户行为路径的深度调研，我们确立了系统的核心功能模块，包括玩家身份自定义、剧情展示与分支选择、角色立绘管理、好感度系统、属性养成、存档读档、AI动态剧情生成、场景切换等。")
replace_paragraph(doc.paragraphs[section_start + 2], "在设计过程中，我们着重关注功能的完整性与扩展性，采用模块化架构设计，确保系统能够轻松适配未来剧情扩展与功能升级。例如，通过JSON格式存储剧情数据，可无缝添加新的剧情分支和角色；利用CSS变量实现主题切换，支持多种古风配色方案；预留AI接口，支持接入不同的大模型服务。此外，系统特别强化了存档功能的实用性，支持6个存档位、自定义存档名和删除功能，满足玩家多存档管理需求。")
replace_paragraph(doc.paragraphs[section_start + 3], "本系统功能设计的核心目标在于打造一个既满足玩家情感体验需求，又能通过技术创新实现高度互动性的文字冒险游戏，为古风游戏开发提供新的技术思路和实践参考。")

# 在section_end之前插入新内容
insert_index = section_end

# 2.3.1 玩家身份自定义功能
doc.paragraphs.insert(insert_index, doc.add_paragraph("2.3.1 玩家身份自定义功能"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("玩家身份自定义功能允许玩家在游戏开始时输入自定义姓名，默认名为'姜嫣'。系统会将玩家姓名存储在全局游戏状态中，并在剧情文本中自动替换。"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("核心代码实现："))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("```javascript"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("let gameState = {"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("    playerName: '姜嫣',"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("    currentNodeId: 'prologue_1',"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("    affection: { xiao: 0, xie: 0, li: 0, chu: 0 },"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("    stats: { wisdom: 50, strength: 30, reputation: 10, perception: 40, beauty: 60, malice: 10 },"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("    flags: {}"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("};"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph(""))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("function startGame() {"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("    const nameInput = document.getElementById('player-name').value.trim();"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("    gameState.playerName = nameInput || '姜嫣';"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("    gameState.currentNodeId = 'prologue_1';"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("    showGameScreen();"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("    renderNode('prologue_1');"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("}"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("```"))
insert_index += 1

# 2.3.2 剧情展示与分支选择功能
doc.paragraphs.insert(insert_index, doc.add_paragraph(""))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("2.3.2 剧情展示与分支选择功能"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("剧情展示功能负责渲染剧情文本、角色名称和选项按钮。分支选择功能支持玩家通过点击选项跳转到不同的剧情节点，实现多分支剧情体验。"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("核心代码实现："))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("```javascript"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("function renderNode(nodeId) {"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("    const node = getNode(nodeId);"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("    gameState.currentNodeId = nodeId;"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("    autoSave();"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph(""))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("    document.getElementById('scene-title').textContent = node.scene || '';"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("    renderCharacters(node.characters);"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph(""))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("    const nameEl = document.getElementById('dialog-name');"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("    const textEl = document.getElementById('dialog-text');"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("    const choicesEl = document.getElementById('dialog-choices');"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph(""))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("    nameEl.textContent = node.character;"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("    textEl.textContent = replacePlayerName(node.text);"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph(""))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("    if (node.choices && node.choices.length > 0) {"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("        node.choices.forEach((choice, index) => {"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("            const btn = document.createElement('button');"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("            btn.className = 'choice-btn';"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("            btn.textContent = replacePlayerName(choice.text);"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("            btn.onclick = () => handleChoice(choice);"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("            choicesEl.appendChild(btn);"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("        });"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("    }"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("}"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("```"))
insert_index += 1

# 2.3.3 角色立绘管理功能
doc.paragraphs.insert(insert_index, doc.add_paragraph(""))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("2.3.3 角色立绘管理功能"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("角色立绘管理功能负责根据剧情节点配置渲染角色立绘，支持左、中、右三个位置的立绘展示。立绘采用透明PNG格式，根据角色脸部朝向自动调整显示位置。"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("核心代码实现："))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("```javascript"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("function renderCharacters(characters) {"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("    const leftEl = document.getElementById('char-left');"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("    const centerEl = document.getElementById('char-center');"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("    const rightEl = document.getElementById('char-right');"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph(""))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("    leftEl.classList.remove('show');"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("    centerEl.classList.remove('show');"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("    rightEl.classList.remove('show');"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph(""))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("    if (characters.left) {"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("        leftEl.src = characters.left;"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("        leftEl.classList.add('show');"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("    }"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("    if (characters.center) {"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("        centerEl.src = characters.center;"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("        centerEl.classList.add('show');"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("    }"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("    if (characters.right) {"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("        rightEl.src = characters.right;"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("        rightEl.classList.add('show');"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("    }"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("}"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("```"))
insert_index += 1

# 2.3.4 好感度系统功能
doc.paragraphs.insert(insert_index, doc.add_paragraph(""))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("2.3.4 好感度系统功能"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("好感度系统功能负责管理四位可攻略男主的好感度数值（0-100），玩家的选择会影响对应男主的好感度。系统提供好感度修改、显示和持久化存储功能。"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("核心代码实现："))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("```javascript"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("function updateAffection(char, amount) {"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("    if (!gameState.affection[char]) {"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("        gameState.affection[char] = 0;"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("    }"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("    gameState.affection[char] = Math.max(0, Math.min(100, gameState.affection[char] + amount));"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph(""))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("    const names = {"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("        xiao: '萧若河',"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("        xie: '谢玉衡',"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("        li: '李自清',"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("        chu: '楚昭云'"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("    };"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph(""))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("    const charName = names[char] || char;"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("    const changeText = amount > 0 ? `+${amount}` : amount;"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("    showAffectionTip(`${charName} 好感度 ${changeText}`);"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("    updateAffectionDisplay();"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("}"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("```"))
insert_index += 1

# 2.3.5 存档读档功能
doc.paragraphs.insert(insert_index, doc.add_paragraph(""))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("2.3.5 存档读档功能"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("存档读档功能支持6个存档位的管理，包括保存、读取、删除和重命名功能。存档数据基于浏览器localStorage实现本地持久化存储，确保游戏进度在重启后不丢失。"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("核心代码实现："))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("```javascript"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("const SAVE_KEY_PREFIX = 'yanluochenyuan_save_';"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("const MAX_SAVE_SLOTS = 6;"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph(""))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("function saveGame(slot) {"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("    const saveData = {"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("        version: '1.2',"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("        playerName: gameState.playerName,"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("        currentNodeId: gameState.currentNodeId,"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("        affection: gameState.affection,"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("        stats: gameState.stats,"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("        flags: gameState.flags,"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("        timestamp: Date.now(),"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("        saveName: '存档 ' + slot"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("    };"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("    localStorage.setItem(SAVE_KEY_PREFIX + slot, JSON.stringify(saveData));"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("    updateSaveSlots();"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("}"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph(""))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("function deleteSave(slot) {"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("    if (confirm('确定要删除此存档吗？')) {"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("        localStorage.removeItem(SAVE_KEY_PREFIX + slot);"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("        updateSaveSlots();"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("    }"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("}"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph(""))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("function renameSave(slot) {"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("    const saveData = localStorage.getItem(SAVE_KEY_PREFIX + slot);"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("    if (saveData) {"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("        const data = JSON.parse(saveData);"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("        const newName = prompt('请输入存档名称：', data.saveName);"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("        if (newName !== null && newName.trim() !== '') {"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("            data.saveName = newName.trim();"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("            localStorage.setItem(SAVE_KEY_PREFIX + slot, JSON.stringify(data));"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("            updateSaveSlots();"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("        }"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("    }"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("}"))
insert_index += 1
doc.paragraphs.insert(insert_index, doc.add_paragraph("```"))

doc.save('D:/嫣落尘渊/孙若妍生产实习报告_新.docx')
print('实验报告已成功生成！')
print(f'总段落数: {len(doc.paragraphs)}')